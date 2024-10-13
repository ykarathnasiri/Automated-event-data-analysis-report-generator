import os
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx2pdf import convert
import pythoncom
import win32com.client

def generate_report():
    pythoncom.CoInitialize()
    # Create necessary folders if they don't exist
    for folder in ['Charts', 'Report']:
        if not os.path.exists(folder):
            os.makedirs(folder)

    # Set paths for saving charts and report
    charts_folder = 'Charts'
    report_folder = 'Report'

    # Read the CSV file
    df = pd.read_csv('Data\event.csv')

    # Preprocess the data
    # Handling Missing Data
    # Drop rows with missing values
    df_cleaned = df.dropna()
    # Fill missing values (customized based on the column type)
    df['Attendee Age'].fillna(df['Attendee Age'].mean(), inplace=True)  # Fill numeric columns with mean
    df['Attendee Contact Information'].fillna('Unknown', inplace=True)  # Fill categorical columns with 'Unknown'
    df['Event Date'] = pd.to_datetime(df['Event Date'], errors='coerce')  # Handle invalid date formats

    # Converting Data Types
    df['Event Date'] = pd.to_datetime(df['Event Date'], format='%Y-%m-%d')  # Convert Event Date to datetime
    df['Ticket Price'] = pd.to_numeric(df['Ticket Price'], errors='coerce')  # Ensure Ticket Price is numeric
    df['Attendee Age'] = pd.to_numeric(df['Attendee Age'], errors='coerce')  # Ensure Age is numeric

    # Handling Duplicates
    df.drop_duplicates(inplace=True)

    # Encoding Categorical Data
    # Convert categorical data into numerical values (for models or specific visualizations)
    df['Attendee Gender'] = df['Attendee Gender'].map({'Male': 1, 'Female': 0})  # Binary encoding for gender

    # Alternatively, you can use one-hot encoding for larger categorical features
    df_encoded = pd.get_dummies(df, columns=['Event Type', 'Ticket Type'], drop_first=True)

    # Creating New Features
    # Extracting year and month from Event Date
    df['Event Year'] = df['Event Date'].dt.year
    df['Event Month'] = df['Event Date'].dt.month

    # Example of a new feature: Age Grouping
    bins = [0, 18, 30, 50, 70, 100]
    labels = ['<18', '18-30', '30-50', '50-70', '70+']
    df['Age Group'] = pd.cut(df['Attendee Age'], bins=bins, labels=labels, include_lowest=True)

    # Outlier Detection and Treatment
    # ORemove outliers using IQR for 'Ticket Price'
    Q1 = df['Ticket Price'].quantile(0.25)
    Q3 = df['Ticket Price'].quantile(0.75)
    IQR = Q3 - Q1
    df = df[~((df['Ticket Price'] < (Q1 - 1.5 * IQR)) | (df['Ticket Price'] > (Q3 + 1.5 * IQR)))]

    # Cap outliers for 'Attendee Age'
    df['Attendee Age'] = df['Attendee Age'].apply(lambda x: 18 if x < 18 else (100 if x > 100 else x))

    # Date-Time Formatting
    # Sorting data by Event Date for time-series visualizations
    df.sort_values(by='Event Date', inplace=True)

    # Visualization
    # 1. Total Ticket Sales by Event
    plt.figure(figsize=(10, 6))
    event_sales = df.groupby('Event Name')['Ticket ID'].count().sort_values(ascending=False)
    event_sales.plot(kind='bar', color=['orange', 'lightblue', 'lightgreen', 'purple', 'pink'])
    plt.title('Total Ticket Sales by Event')
    plt.ylabel('Number of Tickets Sold')
    plt.xlabel('Event Name')
    plt.xticks(rotation=90)
    plt.tight_layout()
    plt.savefig(f'{charts_folder}/total_ticket_sales_by_event.png')
    plt.close()

    # 2. Tickets Sold by Event Type
    plt.figure(figsize=(10, 6))
    event_type_sales = df.groupby('Event Type')['Ticket ID'].count().sort_values(ascending=False)
    event_type_sales.plot(kind='bar', color=['orange', 'lightblue', 'lightgreen', 'purple', 'pink'])
    plt.title('Tickets Sold by Event Type')
    plt.xlabel('Event Type')
    plt.ylabel('Number of Tickets Sold')
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig(f'{charts_folder}/tickets_sold_by_event_type.png')
    plt.close()

    # 3. Event Type Popularity
    event_type_popularity = df['Event Type'].value_counts()
    plt.figure(figsize=(10, 6))
    event_type_popularity.plot(kind='pie', autopct='%1.1f%%', startangle=90,
                               colors=sns.color_palette('coolwarm', len(event_type_popularity)))
    plt.title('Event Type Popularity')
    plt.ylabel('')
    plt.tight_layout()
    plt.savefig(f'{charts_folder}/event_type_popularity.png')
    plt.close()

    # 4. Tickets Sold by Event Organizer
    plt.figure(figsize=(12, 7))
    tickets_by_organizer = df.groupby('Event Organizer')['Ticket ID'].count().sort_values(ascending=False)
    tickets_by_organizer.plot(kind='bar', color=['orange', 'lightblue', 'lightgreen', 'purple', 'pink'])
    plt.title('Total Tickets Sold by Event Organizer')
    plt.xlabel('Event Organizer')
    plt.ylabel('Total Number of Tickets Sold')
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    plt.savefig(f'{charts_folder}/tickets_sold_by_organizer.png')
    plt.close()

    # 5. Ticket Sales Distribution by Price
    plt.figure(figsize=(10, 6))
    sns.histplot(df['Ticket Price'], kde=True, bins=20, color='purple')
    plt.title('Ticket Sales Distribution by Price')
    plt.xlabel('Ticket Price')
    plt.ylabel('No Of Tickets')
    plt.tight_layout()
    plt.savefig(f'{charts_folder}/ticket_sales_distribution_by_price.png')
    plt.close()

    # 6. Average Ticket Price per Event
    avg_ticket_price_per_event = df.groupby('Event Name')['Ticket Price'].mean().sort_values(ascending=False)
    plt.figure(figsize=(10, 6))
    avg_ticket_price_per_event.plot(kind='bar', color=['lightblue', 'lightgreen'])
    plt.title('Average Ticket Price per Event')
    plt.xlabel('Event Name')
    plt.ylabel('Average Ticket Price')
    plt.xticks(rotation=90)
    plt.tight_layout()
    plt.savefig(f'{charts_folder}/avg_ticket_price_per_event.png')
    plt.close()

    # 7. Ticket Type Distribution by Event Type
    fig, axes = plt.subplots(1, 2, figsize=(14, 6))
    # Ticket Type Distribution by Event Type (Bar Chart on the Left)
    sns.countplot(data=df, x='Event Type', hue='Ticket Type', palette='coolwarm', ax=axes[0])
    axes[0].set_title('Ticket Type Distribution by Event Type')
    axes[0].set_xlabel('Event Type')
    axes[0].set_ylabel('Ticket Count')
    axes[0].tick_params(axis='x', rotation=45)
    axes[0].legend(title='Ticket Type')
    #  Ticket Type Popularity (Pie Chart on the Right)
    ticket_type_count = df['Ticket Type'].value_counts()
    ticket_type_count.plot(kind='pie', autopct='%1.1f%%', startangle=90,
                           colors=sns.color_palette('Paired', len(ticket_type_count)), ax=axes[1])
    axes[1].set_title('Ticket Type Popularity')
    axes[1].set_ylabel('')  # Remove the y-label for the pie chart
    axes[1].axis('equal')  # Ensures the pie chart is a circle
    # Adjust layout to avoid overlap
    plt.tight_layout()
    plt.savefig(f'{charts_folder}/ticket_type_distribution.png')
    plt.close()

    # 8. Average Event Duration by Event Type
    avg_event_duration = df.groupby('Event Type')['Event Duration'].mean().sort_values(ascending=False)
    plt.figure(figsize=(10, 6))
    avg_event_duration.plot(kind='bar', color='darkcyan')
    plt.title('Average Event Duration by Event Type')
    plt.xlabel('Event Type')
    plt.ylabel('Average Duration (Hours)')
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig(f'{charts_folder}/avg_event_duration.png')
    plt.close()

    # 9. Attendee Age Distribution by Event Type
    plt.figure(figsize=(10, 6))
    sns.boxplot(data=df, x='Event Type', y='Attendee Age', palette='coolwarm')
    plt.title('Attendee Age Distribution by Event Type')
    plt.xlabel('Event Type')
    plt.ylabel('Attendee Age')
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig(f'{charts_folder}/attendee_age_distribution.png')
    plt.close()

    # 10. Gender Distribution
    fig, axes = plt.subplots(1, 2, figsize=(16, 8))
    # Pie Chart: Overall Attendee Gender Distribution (on the right)
    gender_count = df['Attendee Gender'].value_counts()
    # Assign blue for Male and pink for Female
    gender_count.plot(kind='pie', autopct='%1.1f%%', startangle=90,
                      colors=['lightblue', 'pink'], ax=axes[1])
    axes[1].set_title('Overall Attendee Gender Distribution')
    axes[1].set_ylabel('')  # Remove the y-label for the pie chart
    axes[1].axis('equal')  # Ensures the pie chart is a circle
    # Stacked Bar Chart: Gender Distribution by Event (on the left)
    gender_event_distribution = df.groupby(['Event Name', 'Attendee Gender']).size().unstack()
    # Plot the stacked bar chart with blue for Male and pink for Female
    gender_event_distribution.plot(kind='bar', stacked=True, color=['lightblue', 'pink'], ax=axes[0])
    axes[0].set_title('Attendee Gender Distribution by Event')
    axes[0].set_xlabel('Event Name')
    axes[0].set_ylabel('Number of Attendees')
    axes[0].tick_params(axis='x', rotation=45, labelsize=10)  # Rotate x-ticks for readability
    axes[0].legend(title='Gender', labels=['Male', 'Female'], loc='upper right')
    # Adjust layout to avoid overlap
    plt.tight_layout()
    plt.savefig(f'{charts_folder}/gender_distribution.png')
    plt.close()

    # Create a Word document for the report
    doc = Document()

    # Add title
    doc.add_heading('CeylonEvent Analysis Report', 0)

    # Add notice to the document
    notice_paragraph = doc.add_paragraph()
    notice_run = notice_paragraph.add_run(
        'Notice: This CeylonEvent analysis report is for demonstration purposes only. The analysis utilizes a dummy dataset containing cleaned records from August and September 2024, comprising approximately 4,045 entries.')
    notice_run.font.color.rgb = RGBColor(255, 0, 0)  # Set paragraph text to red color

    # Add a paragraph for the report introduction
    doc.add_paragraph(
        'This report provides a comprehensive analysis of event data, including key insights into event popularity, ticket sales, attendee demographics, and more.')

    # Section 1: Event Performance Analysis

    # 1.1 Total Ticket Sales by Event
    # Add section heading
    doc.add_heading('1. Event Performance Analysis', level=1)
    # Add chart to the Word document
    doc.add_picture(f'{charts_folder}/total_ticket_sales_by_event.png', width=Inches(6))
    # Add a paragraph for the chart details
    top_event = event_sales.idxmax()
    top_event_tickets_sold = event_sales.max()
    doc.add_paragraph(
        f'This bar chart depicts the total number of tickets sold for each event. The event with the highest number of tickets sold is {top_event} with {top_event_tickets_sold} tickets sold, indicating its significant popularity.')

    # 1.2 Tickets Sold by Event Type
    # Add chart to the Word document
    doc.add_heading('1.2 Tickets Sold by Event Type', level=2)
    doc.add_picture(f'{charts_folder}/tickets_sold_by_event_type.png', width=Inches(6))
    # Add a paragraph for the chart details
    top_event_type = event_type_sales.idxmax()
    top_event_type_tickets_sold = event_type_sales.max()
    doc.add_paragraph(
        f'The bar chart presents the number of tickets sold for each event type. {top_event_type} has the highest number of tickets sold, indicating its strong popularity among attendees.')

    # 1.3 Event Type Popularity
    doc.add_heading('1.3 Event Type Popularity', level=2)
    doc.add_picture(f'{charts_folder}/event_type_popularity.png', width=Inches(6))
    most_popular_event_type = event_type_popularity.idxmax()
    most_popular_event_type_percentage = event_type_popularity.max() / event_type_popularity.sum() * 100
    doc.add_paragraph(
        f'The pie chart shows the distribution of event types. {most_popular_event_type} is the most popular, accounting for {most_popular_event_type_percentage:.1f}% of events.')

    # 1.4 Tickets Sold by Event Organizer
    doc.add_heading('1.4 Tickets Sold by Event Organizer', level=2)
    doc.add_picture(f'{charts_folder}/tickets_sold_by_organizer.png', width=Inches(6))
    top_organizer = tickets_by_organizer.idxmax()
    top_organizer_tickets_sold = tickets_by_organizer.max()
    doc.add_paragraph(
        f'The bar chart depicts the total number of tickets sold by each event organizer. {top_organizer} has sold the highest number of tickets.')

    # 1.5 Ticket Sales Distribution by Price
    # Add chart to the Word document
    doc.add_heading('1.5 Ticket Sales Distribution by Price', level=2)
    doc.add_picture(f'{charts_folder}/ticket_sales_distribution_by_price.png', width=Inches(6))
    # Add a paragraph for the chart details
    price_range = f"LKR {df['Ticket Price'].min():.2f} to LKR {df['Ticket Price'].max():.2f}"
    doc.add_paragraph(
        f'The histogram displays the distribution of ticket prices. The majority of tickets are priced between {price_range}. This suggests that the event organizers cater to a broad price range to attract diverse audiences.')

    # 1.6 Average Ticket Price per Event
    # Add chart to the Word document
    doc.add_heading('1.6 Average Ticket Price per Event', level=2)
    doc.add_picture(f'{charts_folder}/avg_ticket_price_per_event.png', width=Inches(6))
    # Add a paragraph for the chart details
    highest_avg_price_event = avg_ticket_price_per_event.idxmax()
    highest_avg_price = avg_ticket_price_per_event.max()
    doc.add_paragraph(
        f'The bar chart shows the average ticket price for each event. {highest_avg_price_event} has the highest average ticket price of LKR {highest_avg_price:.2f}.')

    # 1.7 Ticket Type Distribution by Event Type
    # Add chart to the Word document
    doc.add_heading('1.7 Ticket Type Distribution by Event Type', level=2)
    doc.add_picture(f'{charts_folder}/ticket_type_distribution.png', width=Inches(6))
    # Add a paragraph for the chart details
    doc.add_paragraph(
        'The bar chart on the left shows the distribution of ticket types for each event type. The pie chart on the right shows the overall popularity of each ticket type.')

    # 1.8 Average Event Duration by Event Type
    # Add chart to the Word document
    doc.add_heading('1.8 Average Event Duration by Event Type', level=2)
    doc.add_picture(f'{charts_folder}/avg_event_duration.png', width=Inches(6))
    # Add a paragraph for the chart details
    longest_duration_event_type = avg_event_duration.idxmax()
    longest_duration = avg_event_duration.max()
    doc.add_paragraph(
        f'The bar chart shows the average duration of events for each event type. {longest_duration_event_type} has the longest average duration of {longest_duration:.1f} hours.')

    # Section 2: Attendee Demographics Analysis

    # 2.1 Attendee Age Distribution by Event Type
    # Add section heading
    doc.add_heading('2. Attendee Demographics Analysis', level=1)
    # Add chart to the Word document
    doc.add_heading('2.1 Attendee Age Distribution by Event Type', level=2)
    doc.add_picture(f'{charts_folder}/attendee_age_distribution.png', width=Inches(6))
    # Add a paragraph for the chart details
    doc.add_paragraph(
        'The boxplot shows the distribution of attendee ages for each event type.  You can see the age range and any potential outliers for each event type.')

    # 2.2 Gender Distribution
    # Add chart to the Word document
    doc.add_heading('2.2 Gender Distribution', level=2)
    doc.add_picture(f'{charts_folder}/gender_distribution.png', width=Inches(6))
    # Add a paragraph for the chart details
    male_percentage = gender_count[1] / gender_count.sum() * 100
    female_percentage = gender_count[0] / gender_count.sum() * 100
    doc.add_paragraph(
        f'The pie chart shows the overall gender distribution of attendees. {male_percentage:.1f}% are male, while {female_percentage:.1f}% are female. The stacked bar chart provides a breakdown of gender distribution across different events, allowing for a more detailed analysis of audience composition.')

    ## Section 3: Insights and Recommendations
    # Add section heading
    doc.add_heading('3. Insights and Recommendations', level=1)

    # Calculating each event's total attendees
    event_attendees = df.groupby('Event Name')['Attendee Name'].nunique()

    # Calculating each event's revenue
    event_revenue = df.groupby('Event Name')['Ticket Price'].sum()

    # Total revenue from all ticket sales
    total_revenue = df['Ticket Price'].sum()

    # Most popular event type (based on total attendees)
    popular_event_type = df.groupby('Event Type')['Attendee Name'].nunique().idxmax()
    popular_event_type_attendees = df.groupby('Event Type')['Attendee Name'].nunique().max()
    popular_event_type_revenue = df.groupby('Event Type')['Ticket Price'].sum().max()

    # Most popular event (based on total attendees)
    popular_event = df.groupby('Event Name')['Attendee Name'].nunique().idxmax()
    popular_event_attendees = df.groupby('Event Name')['Attendee Name'].nunique().max()
    popular_event_revenue = df.groupby('Event Name')['Ticket Price'].sum().max()

    # Mapping numerical months to their respective month names
    df['Event Month'] = df['Event Month'].replace({8: 'August', 9: 'September'})

    # Total events by month after replacing numbers with names
    events_by_month = df.groupby('Event Month')['Event Name'].nunique()

    # Event organizer with the most events by event type
    organizer_event_type = df.groupby(['Event Organizer', 'Event Type'])['Event Name'].nunique().idxmax()
    organizer_name = organizer_event_type[0]  # Extracting the organizer name
    event_type = organizer_event_type[1]  # Extracting the event type

    # Most active location (where the most attendees are located)
    active_location = df['Attendee Location'].mode()[0]

    # Ticket type with sold ticket counts
    sold_tickets_by_type = df.groupby('Ticket Type')['Ticket ID'].nunique()

    # Add insights to the Word document
    doc.add_heading('Insights', level=2)

    # Use a list for bullet points
    doc.add_paragraph(f'- Total Revenue from Ticket Sales:  LKR  {total_revenue:.2f}')
    doc.add_paragraph(f'- Most Popular Event Type:  {popular_event_type}')
    doc.add_paragraph(f'- Total Attendees for Most Popular Event Type:  {popular_event_type_attendees}')
    doc.add_paragraph(f'- Total Revenue for Most Popular Event Type:  LKR  {popular_event_type_revenue:.2f}')
    doc.add_paragraph(f'- Most Popular Event:  {popular_event}')
    doc.add_paragraph(f'- Total Attendees for Most Popular Event:  {popular_event_attendees}')
    doc.add_paragraph(f'- Total Revenue for Most Popular Event:  LKR  {popular_event_revenue:.2f}')
    doc.add_paragraph(
        f'- Most Active Month for Events:  {events_by_month.idxmax()} with {events_by_month.max()} events')
    doc.add_paragraph(f'- Event Organizer with the Most Events in the  {event_type}  category:  {organizer_name}')
    doc.add_paragraph(f'- Most Active Location for Attendees:  {active_location}')
    doc.add_paragraph(
        f'- Ticket Type with the Most Sold Tickets:  {sold_tickets_by_type.idxmax()} with {sold_tickets_by_type.max()} tickets sold')

    # Add recommendations to the Word document
    doc.add_heading('Recommendations', level=2)

    # Recommendations with details
    doc.add_paragraph(
        f'- Focus on promoting events in the {popular_event_type} category, as this event type has been the most popular in terms of attendees and revenue.')
    doc.add_paragraph(
        f'- Partner with {organizer_name} to organize more events in the {event_type} category, as they have consistently attracted a large audience.')
    doc.add_paragraph(
        f'- Consider offering different ticket types to cater to a wider range of attendee preferences. The {sold_tickets_by_type.idxmax()} ticket type has been the most popular, but diversifying options can potentially increase sales.')
    doc.add_paragraph(
        f'- Implement strategies to attract attendees from {active_location}. This location has consistently yielded the most attendees, suggesting a strong potential market.')



    docx_path = 'Report/event_data_analysis_report.docx'
    doc.save(docx_path)

    # Convert to PDF
    pdf_path = 'Report/event_data_analysis_report.pdf'
    convert(docx_path, pdf_path)


if __name__ == "__main__":
    generate_report()